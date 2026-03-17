import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import json

def apply_style(run, paragraph, config):
    """
    根据配置为一个Run(文字块)和一个段落应用样式
    """
    if not config:
        return
        
    font_name_zh = config.get("font_name_zh")
    font_name_en = config.get("font_name_en")
    size_pt = config.get("size_pt")
    bold = config.get("bold")
    alignment = config.get("alignment")
    line_spacing_multiple = config.get("line_spacing_multiple")
    line_spacing_exact_pt = config.get("line_spacing_exact_pt")
    space_before_pt = config.get("space_before_pt")
    space_after_pt = config.get("space_after_pt")

    # 对齐方式
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 段前段后间距
    if space_before_pt:
        try: paragraph.paragraph_format.space_before = Pt(float(space_before_pt))
        except: pass
    if space_after_pt:
        try: paragraph.paragraph_format.space_after = Pt(float(space_after_pt))
        except: pass

    # 行距设置 (支持固定值和多倍行距)
    if line_spacing_exact_pt:
        try:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(float(line_spacing_exact_pt))
        except: pass
    elif line_spacing_multiple:
        try:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            paragraph.paragraph_format.line_spacing = float(line_spacing_multiple)
        except: pass

    # 字体字号加粗
    if size_pt:
        run.font.size = Pt(float(size_pt))
    if bold is not None:
        run.bold = bold
    
    # 修改中英文字体组合
    if font_name_en:
        run.font.name = font_name_en
    if font_name_zh:
        if not run.font.name:
            run.font.name = font_name_zh # 回退保护
        # 强制设置东亚字体(中文字体有效)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_zh)

def extract_paragraphs_text(input_path):
    """
    提取Word文档中所有段落的文本，带有索引引用的数组返回
    """
    doc = docx.Document(input_path)
    data = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            data.append({"idx": i, "text": text})
    return data

def add_toc_at_index(doc, paragraph_idx):
    """
    使用 OXML 在指定段落之前插入标准 Word 目录及分页符
    """
    try:
        from docx.oxml.shared import OxmlElement
    except ImportError:
        from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import docx
    
    if paragraph_idx < len(doc.paragraphs):
        target_p = doc.paragraphs[paragraph_idx]
        toc_title_p = target_p.insert_paragraph_before()
        toc_p = target_p.insert_paragraph_before()
        page_break_p = target_p.insert_paragraph_before()
    else:
        toc_title_p = doc.add_paragraph()
        toc_p = doc.add_paragraph()
        page_break_p = doc.add_paragraph()
        
    # 添加目录标题
    run = toc_title_p.add_run("目  录")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    toc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 修改目录对应的格式样式
    from docx.enum.style import WD_STYLE_TYPE
    
    # 在中文 Word 环境下，样式名通常为 "TOC 1" 的本地化形式，如 "目录 1"。我们同时配置中英样式名。
    styles_to_update = [
        ('TOC 1', '目录 1', 'toc 1', True), 
        ('TOC 2', '目录 2', 'toc 2', False), 
        ('TOC 3', '目录 3', 'toc 3', False)
    ]
    for std_name, zh_name, alt_name, is_bold in styles_to_update:
        style = None
        # 尝试寻找该样式
        for name in (std_name, zh_name, alt_name):
            try:
                style = doc.styles[name]
                break
            except KeyError:
                continue
                
        if not style:
            # 如果都没找到，则以标准全名添加一个段落样式
            style = doc.styles.add_style(std_name, WD_STYLE_TYPE.PARAGRAPH)
            
        # 设置：宋体，小四(12磅)
        style.font.name = 'SimSun'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        style._element.rPr.rFonts.set(qn('w:ascii'), 'SimSun')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'SimSun')
        style.font.size = Pt(12)
        style.font.bold = is_bold
    
    # 增加段后间距
    toc_title_p.paragraph_format.space_after = Pt(12)
    
    # 添加目录更新域 (TOC field)
    run2 = toc_p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    # 插入默认的提示文本，以便用户右键更新时能看到区域
    run3 = toc_p.add_run("右键点击此处 -> 更新域 -> 更新整个目录，即可自动生成完整目录")
    run3.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
    run3.font.size = Pt(10)
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    
    # 把结束符加在提示文字后面
    run3._r.append(fldChar3)
    
    # 增加分节符 (替换原来的分页符)，从而可以独立设置正文的页码
    pPr = page_break_p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    type_elm = OxmlElement('w:type')
    type_elm.set(qn('w:val'), 'nextPage')
    sectPr.append(type_elm)
    
    # 拷贝主文档的页面大小等设置以防错位
    try:
        main_sectPr = doc._element.body.xpath('.//w:sectPr')[-1]
        for child in main_sectPr:
            if child.tag not in (qn('w:type'), qn('w:headerReference'), qn('w:footerReference')):
                sectPr.append(child.__deepcopy__(True))
    except (IndexError, Exception):
        pass
    pPr.append(sectPr)
    
    # 获取通过我们在上面刚创建的后面那个 Section（正文内容 Section）
    # 由于 python-docx 的 doc.sections 是实时检测 w:sectPr 计算的，此时 doc.sections 应该会更新
    if len(doc.sections) > 1:
        # 清除目录首页可能自带的页脚/页码
        toc_sec = doc.sections[0]
        toc_sec.footer_is_linked_to_previous = False
        for p in toc_sec.footer.paragraphs:
            p.clear()
            
        main_sec = doc.sections[1]
        main_sec.footer_is_linked_to_previous = False
        
        pgNumType = main_sec._sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            main_sec._sectPr.append(pgNumType)
        pgNumType.set(qn('w:start'), '1')
        
        # 往正文页脚插入页码
        footer = main_sec.footer
        if len(footer.paragraphs) == 0:
            footer.add_paragraph()
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 如果原来有内容先清空
        footer_para.clear()
        run_pg = footer_para.add_run()
        
        fldChar1_pg = OxmlElement('w:fldChar')
        fldChar1_pg.set(qn('w:fldCharType'), 'begin')
        instrText_pg = OxmlElement('w:instrText')
        instrText_pg.set(qn('xml:space'), 'preserve')
        instrText_pg.text = "PAGE"
        fldChar2_pg = OxmlElement('w:fldChar')
        fldChar2_pg.set(qn('w:fldCharType'), 'separate')
        fldChar3_pg = OxmlElement('w:fldChar')
        fldChar3_pg.set(qn('w:fldCharType'), 'end')
        
        run_pg._r.append(fldChar1_pg)
        run_pg._r.append(instrText_pg)
        run_pg._r.append(fldChar2_pg)
        run_pg._r.append(fldChar3_pg)
    
    # 强制让 Word 打开时提示更新域
    try:
        settings = doc.settings.element
        updateFields = OxmlElement('w:updateFields')
        updateFields.set(qn('w:val'), 'true')
        settings.append(updateFields)
    except:
        pass

def set_outline_lvl(paragraph, level):
    """
    设置段落的大纲级别，以便 Word 能将其识别为标题（目录根据此级别生成）
    """
    try:
        from docx.oxml.shared import OxmlElement
    except ImportError:
        from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._element.get_or_add_pPr()
    outlineLvl = pPr.find(qn('w:outlineLvl'))
    if outlineLvl is None:
        outlineLvl = OxmlElement('w:outlineLvl')
        pPr.append(outlineLvl)
    outlineLvl.set(qn('w:val'), str(level))

def update_toc_via_com(docx_path):
    """
    尝试使用 Windows COM 接口自动在后台启动 Word 并计算/更新整个文档的目录和页码域。
    只在 Windows 系统上（并且安装了 Microsoft Word 的机器上）生效。
    """
    import os
    try:
        import win32com.client
    except ImportError:
        print("提示: 未安装 pywin32，无法在生成后立刻渲染页码。请手动在打开 Word 时更新域。")
        return
        
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc_path = os.path.abspath(docx_path)
        
        # 打开刚才保存的文档
        doc = word.Documents.Open(doc_path)
        
        # 强制在这个文档里把 TOC 的内置格式改成对应要求 (Word COM 级别修改，防止被底层或模板覆盖)
        try:
            # wdStyleTOC1 = -19, wdStyleTOC2 = -20, wdStyleTOC3 = -21
            toc_styles = [
                (-19, True),   # TOC1: 加粗
                (-20, False),  # TOC2: 不加粗
                (-21, False),  # TOC3: 不加粗
            ]
            for style_id, is_bold in toc_styles:
                style = doc.Styles(style_id)
                style.Font.Name = "宋体"
                style.Font.NameFarEast = "宋体"
                style.Font.Size = 12  # 小四 = 12磅
                style.Font.Bold = is_bold
                
                # 清理默认可能存在的段前段后间距
                style.ParagraphFormat.SpaceBefore = 0
                style.ParagraphFormat.SpaceAfter = 0
                style.ParagraphFormat.LineSpacingRule = 0 # 单倍行距
        except Exception as style_e:
            print(f"微调TOC样式时跳过: {style_e}")

        # 第一遍：更新段落和页码
        doc.Fields.Update()
        # 第二遍：更新目录本身的内容
        for toc in doc.TablesOfContents:
            toc.Update()
            
        # 第三遍：强制应用我们的新样式到目录的所有段落以防“更新目录”又把它洗掉了
        try:
            for toc in doc.TablesOfContents:
                for p in toc.Range.Paragraphs:
                    try:
                        style_name = str(p.Style.NameLocal).strip()
                        # 在 Word 中，目录的段落样式通常叫 "TOC 1" 或 "目录 1", 尾数为级别
                        is_bold_para = style_name.endswith('1')
                    except:
                        is_bold_para = False
                        
                    p.Range.Font.Name = "宋体"
                    p.Range.Font.NameFarEast = "宋体"
                    p.Range.Font.Size = 12
                    p.Range.Font.Bold = is_bold_para
        except Exception as para_e:
            pass

        doc.Save()
        doc.Close()
        word.Quit()
        print("✨ 已经成功通过后台静默计算了所有的排版页码并填充至目录。")
    except Exception as e:
        print(f"⚠️ 自动更新目录内容时发生异常，但不影响文档: {e}")
        try:
            word.Quit()
        except:
            pass

def process_document(input_path, output_path, format_config, paragraph_types, generate_toc=False):
    """
    读取文档 -> 根据 AI 判断的段落类型 -> 应用格式配置 -> 保存新文件
    """
    print(f"正在加载文档: {input_path}")
    doc = docx.Document(input_path)
    
    print("开始应用样式...")
    toc_insert_idx = 0  # 目录始终插入在文档最前方
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
            
        # 从 AI 分类结果中获取当前段落类型，默认为正文
        ptype = paragraph_types.get(str(i), "body")
        
        # 匹配对应配置（如果没有找到此类型的独特配置，则使用空字典）
        current_config = format_config.get(ptype, {})
        print(f" [{ptype}] -> {text[:15]}...")

        # 为标题设置大纲级别，使得 Word 目录域能够抓取
        if ptype == "heading_1":
            set_outline_lvl(p, 0)
        elif ptype == "heading_2":
            set_outline_lvl(p, 1)
        elif ptype == "heading_3":
            set_outline_lvl(p, 2)

        import re
        is_list = False
        # 识别以序号、分点符号开头的段落，不进行不自然的首行缩进
        if re.match(r'^(\d+[\.\、\)]|[\(（]\d+[\)）]|[①-⑩]|\-|•|·)', p.text.strip()):
            is_list = True

        # 应用缩进 (简单处理为首行缩进 2 字符，这在学术排版非常常见)
        if current_config.get("first_line_indent") and not is_list:
            p.paragraph_format.first_line_indent = Pt(28) # 估算 24-28 pt，约等于2个汉字默认宽度
        elif is_list:
            p.paragraph_format.first_line_indent = Pt(0) # 消除奇怪的首行缩进
            
        # 如果是参考文献，不论是否带有编号，都强制为其设定悬挂缩进（左缩进 + 负首行缩进）
        if ptype == "reference":
            p.paragraph_format.left_indent = Pt(24)   # 约设为2字符的悬挂间距
            p.paragraph_format.first_line_indent = Pt(-24)
            
        # 遍历段落的每一个文字块应用样式
        # 由于在遍历过程中我们可能会修改 p.runs，所以这里要取出原始的 runs 列表副本
        original_runs = list(p.runs)
        for run in original_runs:
            text = run.text
            
            # 正文中的引文上标处理，例如 [1], [1,2], [1-3]
            if ptype != "reference" and re.search(r'\[\d+(?:[,\-]\d+)*\]', text):
                from docx.oxml.ns import qn
                parts = re.split(r'(\[\d+(?:[,\-]\d+)*\])', text)
                if len(parts) > 1:
                    r = run._r
                    parent = r.getparent()
                    if parent is not None:
                        idx = parent.index(r)
                        parent.remove(r)  # 移除原先的 run
                        
                        for part in parts:
                            if not part: continue
                            new_run = p.add_run(part)
                            new_r = new_run._r
                            # 改变插入位置，使其替换原来的 run
                            try:
                                p._p.remove(new_r)
                            except ValueError:
                                pass
                            parent.insert(idx, new_r)
                            idx += 1
                            
                            # 应用基础样式
                            apply_style(new_run, p, current_config)
                            
                            # 若匹配文献格式，将其上标并专门设为 Times New Roman
                            if re.match(r'^\[\d+(?:[,\-]\d+)*\]$', part):
                                new_run.font.superscript = True
                                new_run.font.name = "Times New Roman"
                                new_run._element.rPr.rFonts.set(qn('w:ascii'), "Times New Roman")
                                new_run._element.rPr.rFonts.set(qn('w:hAnsi'), "Times New Roman")
                    continue
            
            # 正常处理不需要切分的 run
            # 去除原有的强格式覆盖
            run.font.name = None
            run.font.size = None
            apply_style(run, p, current_config)

    # 插入目录 (如果选项开启)
    if generate_toc:
        print("正在插入标准Word自动目录...")
        add_toc_at_index(doc, toc_insert_idx)

    print(f"处理完成，正在初步保存文档: {output_path}")
    doc.save(output_path)
    
    if generate_toc:
        print("正在调用后台组件计算并在文档中直接生成真实的页码...")
        update_toc_via_com(output_path)
        
    print("大功告成，保存成功！")

if __name__ == "__main__":
    # 简单测试代码（如果有 test.docx 会被创建/修改）
    dummy_config = {
        "title": {"font_name": "黑体", "size_pt": 22.0, "bold": True, "alignment": "center"},
        "heading": {"font_name": "黑体", "size_pt": 16.0, "bold": True},
        "body": {"font_name": "宋体", "size_pt": 12.0, "line_spacing": 1.5, "first_line_indent": True}
    }
    # process_document("test.docx", "test_formatted.docx", dummy_config)
